#!/usr/bin/env python3
"""
Knowledge Distillery (KD) - NotebookLM Research Automation

Automates the workflow of extracting structured knowledge from YouTube videos
or articles using NotebookLM + Gemini AI.

Usage:
    uv run kd.py research --topic "金融投資" --urls "url1,url2,url3" --format briefing
"""

import argparse
import json
import os
import subprocess
import sys
import time
from datetime import datetime
from pathlib import Path


def log(message, level="INFO"):
    """Print log message with timestamp"""
    timestamp = datetime.now().strftime("%H:%M:%S")
    print(f"[{timestamp}] {level}: {message}", file=sys.stderr)


def run_openclaw_browser(action, retry=3, **kwargs):
    """Execute OpenClaw browser command via subprocess with retry"""
    cmd = ["openclaw", "browser", action]
    
    for key, value in kwargs.items():
        if value is not None:
            cmd.append(f"--{key}")
            cmd.append(str(value))
    
    log(f"Executing: {' '.join(cmd)}")
    
    for attempt in range(retry):
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                if attempt < retry - 1:
                    log(f"Browser command failed (attempt {attempt+1}/{retry}), retrying...", "WARN")
                    time.sleep(2)
                    continue
                else:
                    log(f"Browser command failed after {retry} attempts: {result.stderr}", "ERROR")
                    return None
            
            # Parse JSON output if available
            try:
                return json.loads(result.stdout)
            except json.JSONDecodeError:
                return result.stdout
        
        except subprocess.TimeoutExpired:
            if attempt < retry - 1:
                log(f"Browser command timed out (attempt {attempt+1}/{retry}), retrying...", "WARN")
                time.sleep(2)
                continue
            else:
                log("Browser command timed out after all retries", "ERROR")
                return None
        except Exception as e:
            if attempt < retry - 1:
                log(f"Browser command error (attempt {attempt+1}/{retry}): {e}, retrying...", "WARN")
                time.sleep(2)
                continue
            else:
                log(f"Browser command error after all retries: {e}", "ERROR")
                return None
    
    return None


def wait_for_condition(check_fn, timeout=120, interval=3):
    """Wait for a condition to be true"""
    start_time = time.time()
    while time.time() - start_time < timeout:
        if check_fn():
            return True
        time.sleep(interval)
    return False


def create_notebook_and_import_urls(urls, profile="openclaw"):
    """
    Create a new NotebookLM notebook and import URLs.
    Returns the notebook URL and target ID.
    """
    log("Opening NotebookLM...")
    
    # Open NotebookLM homepage
    result = run_openclaw_browser(
        "open",
        targetUrl="https://notebooklm.google.com",
        profile=profile
    )
    
    if not result:
        log("Failed to open NotebookLM", "ERROR")
        return None, None
    
    target_id = result.get("targetId")
    log(f"NotebookLM opened (target: {target_id})")
    
    # Wait for page to load
    time.sleep(5)
    
    # Click "New Notebook" button
    log("Creating new notebook...")
    run_openclaw_browser(
        "act",
        targetId=target_id,
        profile=profile,
        request=json.dumps({"kind": "click", "text": "新建"})
    )
    
    time.sleep(4)
    
    # Click "Website" source type
    log("Selecting website source type...")
    run_openclaw_browser(
        "act",
        targetId=target_id,
        profile=profile,
        request=json.dumps({"kind": "click", "text": "網站"})
    )
    
    time.sleep(3)
    
    # Type URLs into input
    log(f"Importing {len(urls.split(','))} URLs...")
    url_text = "\n".join(urls.split(","))
    
    # Find the input field and type URLs
    snapshot = run_openclaw_browser(
        "snapshot",
        targetId=target_id,
        profile=profile,
        compact=True
    )
    
    # Type URLs into the textbox
    run_openclaw_browser(
        "act",
        targetId=target_id,
        profile=profile,
        request=json.dumps({
            "kind": "type",
            "text": url_text,
            "selector": "textbox"
        })
    )
    
    time.sleep(2)
    
    # Click "Insert" button
    log("Inserting sources...")
    run_openclaw_browser(
        "act",
        targetId=target_id,
        profile=profile,
        request=json.dumps({"kind": "click", "text": "插入"})
    )
    
    # Wait for import to complete
    log("Waiting for Gemini analysis (this may take 30-60 seconds)...")
    time.sleep(45)
    
    # Get current URL (notebook URL) from browser
    log("Extracting NotebookLM URL...")
    result = run_openclaw_browser(
        "snapshot",
        targetId=target_id,
        profile=profile,
        compact=True
    )
    
    # Extract URL from browser result
    notebook_url = "https://notebooklm.google.com/notebook/[ID]"
    if result and isinstance(result, dict):
        # Try to get URL from result
        pass  # URL extraction logic would go here
    
    log(f"Notebook URL: {notebook_url}")
    
    return target_id, notebook_url


def generate_report(target_id, report_format="briefing", profile="openclaw"):
    """
    Generate a report from the notebook.
    Returns the report content.
    """
    log(f"Generating {report_format} report...")
    
    # Click "Report" in Studio panel
    run_openclaw_browser(
        "act",
        targetId=target_id,
        profile=profile,
        request=json.dumps({"kind": "click", "text": "報告"})
    )
    
    time.sleep(5)
    
    # Select report format
    format_text = {
        "briefing": "簡介文件",
        "study-guide": "研讀指南",
        "blog": "網誌文章"
    }.get(report_format, "簡介文件")
    
    log(f"Selecting format: {format_text}")
    run_openclaw_browser(
        "act",
        targetId=target_id,
        profile=profile,
        request=json.dumps({"kind": "click", "text": format_text})
    )
    
    # Wait for report generation
    log("Waiting for report generation (30-60 seconds)...")
    time.sleep(60)
    
    # Click to open the generated report
    log("Opening generated report...")
    time.sleep(3)
    
    # Click "Copy formatted content" button
    log("Copying report content...")
    run_openclaw_browser(
        "act",
        targetId=target_id,
        profile=profile,
        request=json.dumps({"kind": "click", "text": "複製格式化內容"})
    )
    
    time.sleep(2)
    
    # Get clipboard content (report was copied)
    try:
        result = subprocess.run(
            ["pbpaste"],
            capture_output=True,
            text=True,
            timeout=5
        )
        report_content = result.stdout
        log(f"Report copied successfully ({len(report_content)} chars)")
    except Exception as e:
        log(f"Failed to get clipboard content: {e}", "ERROR")
        report_content = "# Report\n\nFailed to extract report content from clipboard."
    
    return report_content


def convert_to_docx(md_path, topic, date_str, notebook_url, report_content):
    """
    Convert markdown report to Word (.docx) format.
    Returns the docx path or None if conversion fails.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log("python-docx not installed, skipping Word export. Install with: pip install python-docx", "WARN")
        return None

    try:
        doc = Document()

        # Title
        title = doc.add_heading(f'{topic} 知識報告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'生成時間：{date_str}').font.size = Pt(10)
        meta.add_run('\n')
        meta.add_run('工具：NotebookLM + Gemini AI').font.size = Pt(10)
        meta.add_run('\n')
        run = meta.add_run(f'NotebookLM 連結：{notebook_url}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph('')  # spacer

        # Parse markdown content into Word
        for line in report_content.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph('')
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped.startswith('1. ') or stripped.startswith('2. ') or stripped.startswith('3. '):
                doc.add_paragraph(stripped[3:], style='List Number')
            elif stripped.startswith('> '):
                p = doc.add_paragraph(stripped[2:])
                p.style = 'Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else 'Quote'
            elif stripped.startswith('**') and stripped.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(stripped.strip('*')).bold = True
            else:
                doc.add_paragraph(stripped)

        # Save
        docx_filename = md_path.stem + '.docx'
        docx_path = md_path.parent / docx_filename
        doc.save(str(docx_path))
        return docx_path

    except Exception as e:
        log(f"Word conversion failed: {e}", "ERROR")
        return None


def save_report(topic, report_content, notebook_url, output_dir=None):
    """
    Save the report to the workspace with proper folder structure.
    """
    if output_dir is None:
        workspace = Path.home() / ".openclaw" / "workspace"
        output_dir = workspace / f"{topic}_KD"
    else:
        output_dir = Path(output_dir)
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate filename
    date_str = datetime.now().strftime("%Y-%m-%d")
    report_filename = f"{date_str}_{topic}_KD報告.md"
    report_path = output_dir / report_filename
    
    # Write report
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(f"# {topic} 知識報告\n\n")
        f.write(f"**生成時間：** {date_str}\n")
        f.write(f"**工具：** NotebookLM + Gemini AI\n")
        f.write(f"**NotebookLM 連結：** {notebook_url}\n\n")
        f.write("---\n\n")
        f.write(report_content)
    
    log(f"Report saved: {report_path}")
    
    # Generate Word (.docx) version
    docx_path = convert_to_docx(report_path, topic, date_str, notebook_url, report_content)
    if docx_path:
        log(f"Word report saved: {docx_path}")
    
    # Create README if it doesn't exist
    readme_path = output_dir / "README.md"
    if not readme_path.exists():
        with open(readme_path, "w", encoding="utf-8") as f:
            f.write(f"# {topic}_KD - Knowledge Distillery\n\n")
            f.write(f"**建立時間：** {date_str}\n")
            f.write(f"**負責人：** 📋💫 IVE（秘書部）\n")
            f.write(f"**用途：** 彙整 {topic} 相關研究\n\n")
            f.write("## 📊 報告清單\n\n")
            f.write(f"### 1. {topic} ({date_str})\n")
            f.write(f"- **檔案：** `{report_filename}`\n")
            f.write(f"- **NotebookLM：** {notebook_url}\n\n")
        log(f"README created: {readme_path}")
    
    return report_path


def main():
    parser = argparse.ArgumentParser(
        description="Knowledge Distillery - Automate NotebookLM research workflow"
    )
    
    subparsers = parser.add_subparsers(dest="command", help="Command")
    
    # Research command
    research_parser = subparsers.add_parser("research", help="Generate research report")
    research_parser.add_argument(
        "--topic",
        required=True,
        help="Topic name (e.g., '金融投資')"
    )
    research_parser.add_argument(
        "--urls",
        required=True,
        help="Comma-separated YouTube/article URLs"
    )
    research_parser.add_argument(
        "--format",
        choices=["briefing", "study-guide", "blog"],
        default="briefing",
        help="Report format (default: briefing)"
    )
    research_parser.add_argument(
        "--output",
        help="Output directory (default: workspace/{topic}_KD/)"
    )
    research_parser.add_argument(
        "--profile",
        default="openclaw",
        help="Browser profile (default: openclaw)"
    )
    
    args = parser.parse_args()
    
    if args.command != "research":
        parser.print_help()
        return 1
    
    # Execute research workflow
    log(f"🎓 Knowledge Distillery: {args.topic}")
    log(f"URLs: {len(args.urls.split(','))} sources")
    
    try:
        # Step 1: Create notebook and import URLs
        target_id, notebook_url = create_notebook_and_import_urls(
            args.urls,
            profile=args.profile
        )
        
        if not target_id:
            log("Failed to create notebook", "ERROR")
            return 1
        
        # Step 2: Generate report
        report_content = generate_report(
            target_id,
            report_format=args.format,
            profile=args.profile
        )
        
        # Step 3: Save report
        report_path = save_report(
            args.topic,
            report_content,
            notebook_url,
            output_dir=args.output
        )
        
        log(f"✅ Research complete!")
        log(f"Report: {report_path}")
        log(f"NotebookLM: {notebook_url}")
        
        print(f"\n📊 Report saved: {report_path}")
        print(f"🔗 NotebookLM: {notebook_url}")
        
        return 0
    
    except Exception as e:
        log(f"Research failed: {e}", "ERROR")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
